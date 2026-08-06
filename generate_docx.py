import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def create_document():
    doc = docx.Document()

    # Set Margins (1 inch all around)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Styling helper functions
    def set_cell_background(cell, fill_hex):
        tcPr = cell._element.get_or_add_tcPr()
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
        tcPr.append(shd)

    def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
        tcPr = cell._element.get_or_add_tcPr()
        tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
        tcPr.append(tcMar)

    # Document Header Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_title_run = p_title.add_run("ScripVault - Capstone Requirements Comparison & Feature Audit")
    p_title_run.font.name = "Calibri"
    p_title_run.font.size = Pt(22)
    p_title_run.font.bold = True
    p_title_run.font.color.rgb = RGBColor(255, 127, 39) # ScripVault Brand Orange

    p_sub = doc.add_paragraph()
    p_sub_run = p_sub.add_run("upGrad KnowledgeHut Fullstack Development Bootcamp (JavaScript / MERN Stack)")
    p_sub_run.font.name = "Calibri"
    p_sub_run.font.size = Pt(13)
    p_sub_run.font.italic = True
    p_sub_run.font.color.rgb = RGBColor(100, 100, 100)

    p_meta = doc.add_paragraph()
    p_meta_run = p_meta.add_run("Project Name: ScripVault (Stock & Mutual Fund Advisory Platform)\nDate: July 29, 2026 | Author: Varun Akshay")
    p_meta_run.font.name = "Calibri"
    p_meta_run.font.size = Pt(10)
    p_meta_run.font.color.rgb = RGBColor(80, 80, 80)

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # --- SECTION 1: EXECUTIVE SUMMARY ---
    h1 = doc.add_heading(level=1)
    h1_run = h1.add_run("1. Executive Summary & Audit Overview")
    h1_run.font.name = "Calibri"
    h1_run.font.color.rgb = RGBColor(255, 127, 39)

    p_exec = doc.add_paragraph()
    p_exec.paragraph_format.line_spacing = 1.15
    p_exec.paragraph_format.space_after = Pt(8)
    p_exec_run = p_exec.add_run(
        "This document provides a comprehensive, feature-by-feature evaluation of the current ScripVault "
        "codebase against the official upGrad KnowledgeHut Capstone Project Guidebook requirements for the "
        "Stock Market Project track.\n\n"
        "Overall Project Completion Status: ~85% COMPLETE\n"
        "Core authentication, portfolio management, dashboard summaries, explore scrips dataset, watchlist, and expert advisory query submission "
        "are fully functional. Key remaining tasks include implementing real-time price streaming via WebSockets/SSE for stock charts and configuring Express static asset serving for monolithic production deployment."
    )
    p_exec_run.font.name = "Calibri"
    p_exec_run.font.size = Pt(11)

    # --- SECTION 2: COMPARISON CHECKLIST TABLE ---
    h2 = doc.add_heading(level=1)
    h2_run = h2.add_run("2. Institution Requirements Comparison Checklist")
    h2_run.font.name = "Calibri"
    h2_run.font.color.rgb = RGBColor(255, 127, 39)

    table_data = [
        ["Req ID", "Requirement Category", "Institution Specification (upGrad)", "Current Implementation Status", "Action Plan / Status Note"],
        
        ["REQ-01", "User Accounts & Auth", "Visitors can sign up and login with email ID & password. JWT authorization for protected routes.", "✅ COMPLETED (100%)", "JWT login & signup implemented. Lowercase & trim normalization added in server/routes/auth.js."],
        
        ["REQ-02", "User Dashboard - Net Worth", "Display user's net worth based on investments held (arbitrary or dynamic figure).", "✅ COMPLETED (100%)", "Net worth summary cards and chart implemented in Dashboard.js."],
        
        ["REQ-03", "User Dashboard - Market & Invested Value", "Display total Market Value and total Invested Value of purchased stocks on dashboard.", "⚠️ PARTIALLY DONE (85%)", "Basic investment metrics display; enhancement needed to explicitly show Market Value vs. Invested Value totals."],
        
        ["REQ-04", "Portfolio Management", "Categorized list of stocks & mutual funds held. Support One-time and SIP investment flows.", "✅ COMPLETED (100%)", "Portfolio.js & Investments.js handle holdings list and One-time/SIP investment additions."],
        
        ["REQ-05", "Profile Management", "Allow users to update phone number, password, and postal address details only.", "✅ COMPLETED (100%)", "Profile.js & profile.js API enable editing phone, password, and address details."],
        
        ["REQ-06", "Explore Scrips & Fake API", "Search and add scrips/stocks. Express backend API produces fake dataset of scrips.", "✅ COMPLETED (100%)", "Explore.js & GET /api/explore serve 15+ pre-seeded stocks, mutual funds, ETFs, and NFOs."],
        
        ["REQ-07", "Financial Categories", "Home/Explore selection for Mutual funds, Fixed deposits, ETFs, NFOs, NPS.", "✅ COMPLETED (90%)", "Filter tabs for Mutual Funds, Stocks, ETFs, NFOs, FDs, NPS integrated on Explore page."],
        
        ["REQ-08", "Real-Time Price Streaming & Charts", "Provide chart of stock with Buy/Sell buttons. Use WebSockets or SSE for random price updates (1y, 3y, 5y returns).", "❌ YET TO BE DONE (30%)", "High Priority: Implement WebSockets (Socket.IO) or SSE endpoint for real-time updating stock quotes and live chart updates."],
        
        ["REQ-09", "Watchlist Feature", "Mark scrips from Explore as 'Watchlist' and display on dashboard/watchlist view.", "✅ COMPLETED (100%)", "Watchlist.js & watchlist.js API allow adding, viewing, and deleting saved scrips."],
        
        ["REQ-10", "Ask Experts Advisory", "Suggest portfolios based on needs (long-term, short-term). Log queries for backend/admin view.", "✅ COMPLETED (85%)", "AskExperts.js & POST /api/ask-experts/submit-query allow submitting queries with expert responses."],
        
        ["REQ-11", "Monolithic Production Serving", "Express backend must serve static frontend React build (`/build`) on root `/` route.", "⏳ PENDING SETUP (50%)", "Medium Priority: Configure express.static in server/index.js to serve react build for monolithic deployment."]
    ]

    table = doc.add_table(rows=len(table_data), cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_cells = table.rows[0].cells
    for i, title in enumerate(table_data[0]):
        hdr_cells[i].text = title
        set_cell_background(hdr_cells[i], "FF7F27")
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        hdr_cells[i].paragraphs[0].runs[0].font.name = "Calibri"

    col_widths = [Inches(0.7), Inches(1.5), Inches(2.2), Inches(1.4), Inches(2.2)]

    for row_idx, row in enumerate(table.rows[1:], start=1):
        data = table_data[row_idx]
        bg_color = "F9F9F9" if row_idx % 2 == 0 else "FFFFFF"
        for col_idx, cell in enumerate(row.cells):
            cell.text = data[col_idx]
            set_cell_background(cell, bg_color)
            set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
            p = cell.paragraphs[0]
            p.paragraph_format.line_spacing = 1.05
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            if len(p.runs) > 0:
                p.runs[0].font.name = "Calibri"
                p.runs[0].font.size = Pt(9.5)
                if "✅" in data[col_idx]:
                    p.runs[0].font.color.rgb = RGBColor(0, 128, 0)
                elif "⚠️" in data[col_idx] or "⏳" in data[col_idx]:
                    p.runs[0].font.color.rgb = RGBColor(180, 100, 0)
                elif "❌" in data[col_idx]:
                    p.runs[0].font.color.rgb = RGBColor(200, 0, 0)

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # --- SECTION 3: DETAILED ANALYSIS & NEXT STEPS ---
    h3 = doc.add_heading(level=1)
    h3_run = h3.add_run("3. Detailed Feature Breakdown & Action Items")
    h3_run.font.name = "Calibri"
    h3_run.font.color.rgb = RGBColor(255, 127, 39)

    items = [
        ("1. Real-Time Price Streaming (SSE / WebSockets)", 
         "Institution Requirement: Page 13 states that while exploring stocks, the app must provide real-time price updates using WebSockets or SSE for 1y, 3y, and 5y return charts with Buy/Sell actions.\n"
         "Action Required: Add a Server-Sent Events (SSE) or Socket.IO module in Express backend to broadcast price changes every 3-5 seconds, and update Explore.js to render live chart updates with Buy/Sell triggers."),

        ("2. Express Monolithic Static Build Serving",
         "Institution Requirement: Page 5 & 7 state that the Express backend must expose APIs on /api while serving the React production build (`scripvault-frontend/build`) on the root `/` route for simple monolith deployment.\n"
         "Action Required: Add `npm run build` step to package.json and configure `express.static(path.join(__dirname, '../scripvault-frontend/build'))` in `server/index.js`."),

        ("3. Dashboard Market Value vs. Invested Value Enhancement",
         "Institution Requirement: Page 13 specifies that the Dashboard should explicitly show Market Value of stock purchased and Invested Value of stock purchased.\n"
         "Action Required: Enhance `Dashboard.js` summary metrics to display explicit side-by-side 'Invested Value' vs 'Current Market Value' breakdown cards calculated from user portfolio holdings."),

        ("4. Deployment & GitHub Submission Preparation",
         "Institution Requirement: Page 8 specifies that a clean GitHub repository and deployment URL must be ready for submission.\n"
         "Action Required: Verify environment variables, update README.md with comprehensive installation and API endpoint tables, and confirm single-command startup.")
    ]

    for title, desc in items:
        p_item = doc.add_paragraph()
        p_item.paragraph_format.space_after = Pt(6)
        r_title = p_item.add_run(title + "\n")
        r_title.bold = True
        r_title.font.name = "Calibri"
        r_title.font.size = Pt(11)
        r_title.font.color.rgb = RGBColor(50, 50, 50)
        
        r_desc = p_item.add_run(desc)
        r_desc.font.name = "Calibri"
        r_desc.font.size = Pt(10)

    # Save document
    doc.save("ScripVault_Requirements_Comparison_Checklist.docx")
    print("ScripVault_Requirements_Comparison_Checklist.docx generated successfully!")

if __name__ == "__main__":
    create_document()
